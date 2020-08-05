#Author: Phuc Tran

import pdfplumber
import re
from pdf2docx.main import parse
from docx import *

#get the first page and convert it to docx because the first page often contains title
def pdf2word(pdf_path):
	docx_file = 'firstpage.docx'
	#get the first page
	parse(pdf_path, docx_file, start=0, end=1) 
	return docx_file


#get the big word because the title is often bigger than the orther words
def getBigWords(pdf_path):
	document = Document(pdf2word(pdf_path))
	size_list = []
	for para in document.paragraphs:
		for run in para.runs:
			#the font obliges '!=None'
			if (run.font.size != None):
				#append it to a list name size_list
				size_list.append(run.font.size)
	#get max size from list
	size_max = max(size_list)
	get_title = []
	for para in document.paragraphs:
			for run in para.runs:
				#choose the words have biggest size as title
				if (run.font.size == size_max):
					get_title.append(run.text)
	return get_title
	

#get bold text because the title often have bold style
def getBoldWords(pdf_path):
	document = Document(pdf2word(pdf_path))
	bolds_list=[]
	for para in document.paragraphs:
	    for run in para.runs:
	        if run.bold:
	            bolds_list.append(run.text)
	return bolds_list

#find a list of words that is consecutive uppercase 
def findTitleByIndexConsecutiveUpperCaseWords():
	index_word_list = []
	for word in word_list:
		if word.isupper():
			ind = word_list.index(word)
			index_word_list.append(ind)
	length = len(index_word_list)

	title_index_list = []
	for i in range(0, length - 1):
		for j in range(1, length):
			if ((index_word_list[i] + 1) == index_word_list[j]):
				title_index_list.append(index_word_list[i])
				title_index_list.append(index_word_list[j])
	title_index_list = list(dict.fromkeys(title_index_list))
	return title_index_list

#get title from condition 3(find Title By Consecutive Upper Case Words)
def getTitleFromUpperCaseWords(word_list, title_index_list):
	title =''
	for i in title_index_list:
		title = title + word_list[i] + ' '
	return title


#get title from condition 2(find Title By Bold Words)
def getTitleFromBoldWords(pdf_path):
	title = ''
	w = getBoldWords(pdf_path)
	for word in w:
		if word.isupper():
			title = title + word + ' '
	return title


#get title from condition 1(find Title By Biggest Words)
def getTitleFromBiggestWords(pdf_path):
	title = ' '
	w = getBigWords(pdf_path)
	for word in w:
		title = title + word + ' '
	return title

#select the title is longest
def finalTitle(c1, c2, c3):
	title = []
	title.append(c1)
	title.append(c2)
	title.append(c3)

	return max(title)

if __name__ == '__main__':

	#Read file pdf
	pdf_path = 'Notice SFA 02-N02 Supervision of Market Participants.pdf' #file path here
	pdf = pdfplumber.open(pdf_path) 

	#get first page include title
	page = pdf.pages[0] 
	text = page.extract_text()
	#split word from pdf file
	word_list = re.sub("[^\w]", " ",  text).split()	
	
	condition1 = getTitleFromBiggestWords(pdf_path)
	condition2 = getTitleFromBoldWords(pdf_path)

	title_index_list = findTitleByIndexConsecutiveUpperCaseWords()
	condition3 = getTitleFromUpperCaseWords(word_list, title_index_list)

	print('The title is: ', finalTitle(condition1, condition2, condition3))

	'''
	print('Condition 1st:')
	print(getTitleFromBiggestWords(pdf_path))
	print('\n')

	print('Condition 2nd:')
	print(getTitleFromBoldWords(pdf_path))
	print('\n')

	print('Condition 3rd:')
	title_index_list = findTitleByIndexConsecutiveUpperCaseWords()
	print(getTitleFromUpperCaseWords(word_list, title_index_list))

	'''